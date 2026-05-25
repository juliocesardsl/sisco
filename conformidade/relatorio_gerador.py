"""
Gerador de relatórios em formato Word para divergências de carga horária.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from io import BytesIO


def gerar_relatorio_carga_horaria(resultados, rubrica, ano, carga_cadastrada, mes):
    """
    Gera um documento Word com relatórios individuais para cada servidor com divergência de carga horária.
    
    Args:
        resultados: Lista de resultados de verificação (apenas com carga horária divergente)
        rubrica: Código da rubrica
        ano: Ano da referência
        carga_cadastrada: Carga horária cadastrada (a esperada)
        mes: Mês da referência
    
    Returns:
        BytesIO: Documento Word em memória pronto para download
    """
    doc = Document()
    
    # Configurar espaçamento padrão
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Formatar data atual em português
    data_atual = datetime.now().strftime('%d de %B de %Y').replace(
        'January', 'janeiro'
    ).replace(
        'February', 'fevereiro'
    ).replace(
        'March', 'março'
    ).replace(
        'April', 'abril'
    ).replace(
        'May', 'maio'
    ).replace(
        'June', 'junho'
    ).replace(
        'July', 'julho'
    ).replace(
        'August', 'agosto'
    ).replace(
        'September', 'setembro'
    ).replace(
        'October', 'outubro'
    ).replace(
        'November', 'novembro'
    ).replace(
        'December', 'dezembro'
    )

    # Normalizar mês de referência para exibição
    mes_nome = ''
    if isinstance(mes, int):
        mes_num = mes
    elif isinstance(mes, str) and mes.isdigit():
        mes_num = int(mes)
    else:
        mes_num = None

    meses_portugues = {
        1: 'janeiro',
        2: 'fevereiro',
        3: 'março',
        4: 'abril',
        5: 'maio',
        6: 'junho',
        7: 'julho',
        8: 'agosto',
        9: 'setembro',
        10: 'outubro',
        11: 'novembro',
        12: 'dezembro',
    }
    if mes_num in meses_portugues:
        mes_nome = meses_portugues[mes_num]
    elif isinstance(mes, str) and mes:
        mes_nome = mes
    else:
        mes_nome = 'mês não informado'

    # Gerar um relatório para cada servidor
    for idx, resultado in enumerate(resultados):
        # Título
        # titulo = doc.add_paragraph()
        # titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # run = titulo.add_run('DIREÇÃO DE CONFORMIDADE DA FOLHA DE PAGAMENTO')
        # run.font.size = Pt(12)
        # run.font.bold = True
        
        # # Subtítulo
        # subtitulo = doc.add_paragraph()
        # subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # run = subtitulo.add_run('(DICOFP)')
        # run.font.size = Pt(11)
        # run.font.bold = True
        
        # Assunto
        assunto = doc.add_paragraph()
        assunto.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = assunto.add_run('Assunto: ')
        run.font.bold = True
        run = assunto.add_run('Desconformidade no regime de trabalho.')
        
        # Data
        doc.add_paragraph()
        
        # Dados do servidor
        nome_servidor = resultado.get('nome_servidor', '')
        matricula = resultado.get('matricula', '')
        empresa = resultado.get('empresa', '')
        nome_do_orgao = resultado.get('orgao', '')
        if empresa and nome_do_orgao:
            empresa_orgao = f"{nome_do_orgao}"
        elif nome_do_orgao:
            empresa_orgao = nome_do_orgao
        else:
            empresa_orgao = empresa
        
        # Conteúdo principal
        intro = doc.add_paragraph(
            f'A Diretoria de Conformidade da Folha de Pagamento (DICOFP), no exercício das '
            f'atribuições previstas no art. 215 da Portaria nº 544, de 11 de julho de 2025, realizou '
            f'análise dos registros de pagamento relacionados à carga horária cadastrada dos '
            f'servidores, com o objetivo de identificar eventuais desconformidades cadastrais.'
        )
        
        doc.add_paragraph(
            f'Nesse contexto, verificou-se, durante a análise da folha de pagamento referente ao mês de '
            f'{mes_nome} de {ano}, que o servidor(a) {nome_servidor}, matrícula nº {matricula}, '
            f'encontra-se cadastrado(a) no Sistema Único de Gestão de Recursos Humanos (SIGRH) com carga horária de '
            f'{carga_cadastrada}h, divergente da informação registrada no SIGRWEB.'
        )
        
        doc.add_paragraph(
            f'Diante do exposto, submete-se o presente processo à Subsecretaria de Administração da '
            f'Folha de Pagamento (SUAFP), com sugestão de encaminhamento ao(à) {empresa_orgao}, '
            f'para que informe qual é a carga horária correta do(a) referido(a) servidor(a) e, caso necessário, '
            f'promova a devida regularização cadastral nos sistemas corporativos, a fim de assegurar a conformidade das informações '
            f'e a uniformização dos registros entre os sistemas SIGRH e SIGRWEB.'
        )
        
        # Rodapé
        # doc.add_paragraph()
        # assinatura = doc.add_paragraph()
        # assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # run = assinatura.add_run('_' * 40)
        # run.font.size = Pt(11)
        
        # responsavel = doc.add_paragraph()
        # responsavel.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # run = responsavel.add_run('Responsável pela Conformidade')
        # run.font.size = Pt(10)
        
        # Quebra de página entre relatórios (exceto no último)
        if idx < len(resultados) - 1:
            doc.add_page_break()
    
    # Salvar em memória
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io
